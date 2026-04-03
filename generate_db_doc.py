#!/usr/bin/env python3
"""生成秦人户外小程序数据库说明文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ===== 样式设置 =====
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ===== 封面标题 =====
title = doc.add_heading('秦人户外小程序', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('云端数据库说明文档', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# 文档信息
info_table = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
info_data = [
    ('项目名称', '秦人户外集合 - 陕西户外徒步路线查询小程序'),
    ('云环境ID', 'cloud1-1ghoxvn859e9d0df'),
    ('AppID', 'wx45a27eae361c8cb5'),
    ('数据库类型', '微信云开发 CloudBase（MongoDB）'),
    ('文档生成日期', '2026-04-03'),
]
for i, (k, v) in enumerate(info_data):
    info_table.rows[i].cells[0].text = k
    info_table.rows[i].cells[1].text = v

doc.add_page_break()

# ===== 目录概览 =====
doc.add_heading('目录概览', level=1)
doc.add_paragraph('本项目共使用 7 个云端数据集合：')

collections_overview = [
    ('routes', '徒步路线', '存储所有徒步路线的详细信息（约50条）'),
    ('users', '用户信息', '存储用户注册信息、昵称、访问次数等'),
    ('user_data', '用户数据', '存储用户的收藏、已走过记录、装备清单等'),
    ('articles', '知识文章', '存储户外知识文章（装备、安全、礼仪等）'),
    ('admin_config', '后台配置', '存储后台管理系统配置信息'),
    ('counters', '计数器', '存储自增序列（如用户编号）'),
    ('weather', '天气缓存', '缓存天气数据，定时刷新'),
]

overview_table = doc.add_table(rows=len(collections_overview)+1, cols=3, style='Light Grid Accent 1')
overview_table.rows[0].cells[0].text = '集合名称'
overview_table.rows[0].cells[1].text = '中文名'
overview_table.rows[0].cells[2].text = '说明'
for i, (name, cn, desc) in enumerate(collections_overview):
    overview_table.rows[i+1].cells[0].text = name
    overview_table.rows[i+1].cells[1].text = cn
    overview_table.rows[i+1].cells[2].text = desc

doc.add_paragraph('')


# ===== 辅助函数：创建字段表 =====
def add_field_table(doc, fields):
    """fields: list of (字段名, 类型, 中文名, 说明, 是否必填, 默认值)"""
    table = doc.add_table(rows=len(fields)+1, cols=6, style='Light Grid Accent 1')
    headers = ['字段名', '类型', '中文名', '说明', '必填', '默认值']
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for i, row in enumerate(fields):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = str(val)
    doc.add_paragraph('')


# ===== 1. routes 集合 =====
doc.add_page_break()
doc.add_heading('1. routes（徒步路线）', level=1)

doc.add_heading('1.1 集合概述', level=2)
doc.add_paragraph('存储所有徒步路线的详细信息，包括路线基本信息、难度、风景、分段描述、装备建议、安全提示等。这是小程序最核心的数据集合。')

doc.add_heading('1.2 字段结构', level=2)
add_field_table(doc, [
    ('_id', 'String', '路线ID', '路线唯一标识，格式 route_XXX', '是', '系统生成'),
    ('name', 'String', '路线名称', '路线全名（含区域标注）', '是', '-'),
    ('description', 'String', '路线简介', '路线简要描述', '是', '-'),
    ('coverImage', 'String', '封面图路径', '封面图URL或本地路径', '否', '""'),
    ('images', 'Array[String]', '图片列表', '路线相关图片URL数组', '否', '[]'),
    ('difficulty.level', 'Number', '难度等级', '1-5数字，1最简单', '是', '-'),
    ('difficulty.label', 'String', '难度标签', '如：第一次也能走、稍微有点挑战', '是', '-'),
    ('difficulty.suitableFor', 'Array[String]', '适合人群', '如：新手、亲子5岁+、老年人', '否', '[]'),
    ('distance_km', 'Number', '距离(公里)', '路线总距离', '是', '-'),
    ('duration_hours', 'Number', '预计时长(小时)', '完成路线预计时间', '是', '-'),
    ('elevation_gain_m', 'Number', '累计爬升(米)', '路线总爬升高度', '否', '0'),
    ('cost.type', 'String', '费用类型', '免费/收费', '否', '"免费"'),
    ('cost.amount', 'Number', '费用金额', '费用数额（元）', '否', '0'),
    ('cost.note', 'String', '费用说明', '费用备注', '否', '""'),
    ('scenery', 'Array[String]', '风景标签', '如：溪流、瀑布、红叶、森林', '否', '[]'),
    ('location.direction', 'String', '方位方向', '秦岭东线/中线/西线', '是', '-'),
    ('location.address', 'String', '详细地址', '路线所在位置描述', '是', '-'),
    ('location.navAddress', 'String', '导航地址', '导航搜索关键字', '否', '""'),
    ('location.publicTransport', 'String', '交通方式', '自驾/公交/地铁等交通信息', '否', '""'),
    ('sections', 'Array[Object]', '路段分段', '路线分段详细描述', '否', '[]'),
    ('sections[].name', 'String', '分段名称', '如：0-5km', '是', '-'),
    ('sections[].road', 'String', '路面类型', '如：水泥路、山间小道、石阶', '否', '""'),
    ('sections[].desc', 'String', '分段描述', '该段路况和注意事项', '否', '""'),
    ('equipment.must', 'Array[String]', '必备装备', '必须携带的装备清单', '否', '[]'),
    ('equipment.suggest', 'Array[String]', '建议装备', '建议携带的装备', '否', '[]'),
    ('equipment.noNeed', 'Array[String]', '无需携带', '不需要带的物品', '否', '[]'),
    ('safety.warnings', 'Array[String]', '安全警告', '安全注意事项列表', '否', '[]'),
    ('safety.emergencyPhone', 'String', '紧急电话', '救援电话信息', '否', '""'),
    ('best_season', 'String', '最佳季节', '春/夏/秋/冬/全年', '否', '"全年"'),
    ('latitude', 'Number', '纬度', '路线起点纬度', '否', 'null'),
    ('longitude', 'Number', '经度', '路线起点经度', '否', 'null'),
    ('order', 'Number', '排序号', '展示顺序，越小越靠前', '否', '999'),
    ('isActive', 'Boolean', '是否启用', '控制路线是否在前台展示', '否', 'true'),
    ('favoriteCount', 'Number', '收藏次数', '用户收藏该路线的总次数', '否', '0'),
    ('completedCount', 'Number', '已完成次数', '用户标记已走过的总次数', '否', '0'),
    ('highlights', 'String', '路线亮点', '路线的精华亮点描述', '否', '""'),
    ('createdAt', 'String/Date', '创建时间', '数据创建时间', '否', '系统自动'),
])

doc.add_heading('1.3 索引信息', level=2)
doc.add_paragraph('• order 字段：按升序排列，用于列表默认排序')
doc.add_paragraph('• isActive 字段：用于筛选有效路线')
doc.add_paragraph('• location.direction 字段：用于按方向（东/中/西线）筛选')
doc.add_paragraph('• difficulty.level 字段：用于按难度筛选')
doc.add_paragraph('• favoriteCount / completedCount：用于排行榜聚合查询')

doc.add_heading('1.4 读写权限', level=2)
doc.add_paragraph('• 前端：仅读（通过 routes 云函数查询列表/详情/搜索）')
doc.add_paragraph('• 云函数：完整读写（routes、user-data、admin-api、import-data、init-stats、highlights-update 云函数均可操作）')

doc.add_heading('1.5 典型数据示例', level=2)
doc.add_paragraph('''{
  "_id": "route_001",
  "name": "临潼骊山最美环山公路骑行徒步（临潼区）",
  "description": "骊山环山公路被誉为\"西安最美盘山路\"",
  "difficulty": { "level": 1, "label": "第一次也能走", "suitableFor": ["新手", "亲子5岁+", "老年人"] },
  "distance_km": 15,
  "duration_hours": 4,
  "elevation_gain_m": 300,
  "cost": { "type": "免费", "amount": 0, "note": "环山路" },
  "scenery": ["盘山公路", "日落晚霞", "俯瞰城市", "山花烂漫"],
  "location": { "direction": "秦岭东线", "address": "陕西省西安市临潼区骊山周边" },
  "sections": [{ "name": "0-5km", "road": "水泥路/土路", "desc": "平缓起步，适合热身" }],
  "equipment": { "must": ["徒步鞋", "零食干粮", "水(至少2L)"], "suggest": ["防晒帽"] },
  "order": 1,
  "isActive": true,
  "favoriteCount": 0,
  "completedCount": 0
}''')


# ===== 2. users 集合 =====
doc.add_page_break()
doc.add_heading('2. users（用户信息）', level=1)

doc.add_heading('2.1 集合概述', level=2)
doc.add_paragraph('存储用户注册信息。用户首次进入小程序时由 user-data 云函数的 init-user 操作创建，记录用户编号、昵称和访问次数等基础信息。')

doc.add_heading('2.2 字段结构', level=2)
add_field_table(doc, [
    ('_id', 'String', '文档ID', '系统自动生成', '是', '系统生成'),
    ('_openid', 'String', '微信OpenID', '用户微信唯一标识', '是', '微信自动注入'),
    ('userNumber', 'Number', '用户编号', '自增用户编号，从counters集合获取', '是', '自动分配'),
    ('nickName', 'String', '昵称', '默认格式：XXX号徒步爱好者', '否', '"XXX号徒步爱好者"'),
    ('avatarUrl', 'String', '头像URL', '用户头像', '否', '""'),
    ('visitCount', 'Number', '访问次数', '用户累计访问小程序次数', '否', '1'),
    ('createdAt', 'Date', '注册时间', '用户首次注册时间', '否', 'db.serverDate()'),
])

doc.add_heading('2.3 索引信息', level=2)
doc.add_paragraph('• _openid 字段：按用户OpenID查询，每个OpenID唯一对应一个用户')
doc.add_paragraph('• userNumber 字段：自增编号，用于展示"第X号徒步爱好者"')

doc.add_heading('2.4 读写权限', level=2)
doc.add_paragraph('• 前端：通过 user-data 云函数间接读写')
doc.add_paragraph('• 云函数：user-data（创建/更新）、admin-api（管理查询/列表导出）')

doc.add_heading('2.5 典型数据示例', level=2)
doc.add_paragraph('''{
  "_id": "auto_generated_id",
  "_openid": "o4BVT3QcW1wbAgEt1yqRD7drVPhY",
  "userNumber": 42,
  "nickName": "042号徒步爱好者",
  "avatarUrl": "",
  "visitCount": 5,
  "createdAt": "2026-03-29T12:00:00.000Z"
}''')


# ===== 3. user_data 集合 =====
doc.add_page_break()
doc.add_heading('3. user_data（用户数据）', level=1)

doc.add_heading('3.1 集合概述', level=2)
doc.add_paragraph('存储用户的行为数据，包括收藏路线、已走过记录和装备清单勾选状态。每条记录对应一个用户（按 _openid 关联）。这是功能最丰富的集合。')

doc.add_heading('3.2 字段结构', level=2)
add_field_table(doc, [
    ('_id', 'String', '文档ID', '系统自动生成', '是', '系统生成'),
    ('_openid', 'String', '微信OpenID', '用户微信唯一标识', '是', '微信自动注入'),
    ('favorites', 'Array[Object]', '收藏列表', '用户收藏的路线列表', '否', '[]'),
    ('favorites[].routeId', 'String', '路线ID', '收藏的路线ID（关联routes集合）', '是', '-'),
    ('favorites[].date', 'String', '收藏日期', 'ISO日期字符串', '是', '-'),
    ('completed', 'Array[Object]', '已走过记录', '用户标记的已走过路线记录', '否', '[]'),
    ('completed[].routeId', 'String', '路线ID', '已完成的路线ID（关联routes集合）', '是', '-'),
    ('completed[].date', 'String', '完成日期', '完成日期 YYYY-MM-DD 格式', '是', '-'),
    ('completed[].name', 'String', '路线名称', '路线名称（冗余存储）', '否', '""'),
    ('completed[].weather', 'String', '天气', '当天天气情况', '否', '""'),
    ('completed[].feeling', 'String', '感受', '用户感受描述', '否', '""'),
    ('completed[].difficultyFeeling', 'String', '难度感受', '用户对难度的评价', '否', '""'),
    ('completed[].companions', 'String', '同行人', '同行伙伴描述', '否', '""'),
    ('completed[].distance', 'Number', '实际距离', '实际行走距离（公里）', '否', '0'),
    ('completed[].note', 'String', '备注', '用户备注信息', '否', '""'),
    ('completed[].completedAt', 'Number', '完成时间戳', '毫秒级时间戳，用于去重和排序', '是', 'Date.now()'),
    ('checklists', 'Object', '装备清单', '按路线ID组织的装备勾选状态', '否', '{}'),
    ('checklists[routeId].checked', 'Array', '已勾选项', '用户勾选的装备项索引', '否', '[]'),
    ('checklists[routeId].custom', 'Array', '自定义项', '用户添加的自定义装备', '否', '[]'),
    ('updatedAt', 'Date', '更新时间', '最后更新时间', '否', 'db.serverDate()'),
])

doc.add_heading('3.3 索引信息', level=2)
doc.add_paragraph('• _openid 字段：按用户OpenID查询，一个用户最多一条记录')
doc.add_paragraph('• completed[].routeId + completed[].date：联合唯一，用于防止同一天重复标记同一路线')

doc.add_heading('3.4 读写权限', level=2)
doc.add_paragraph('• 前端：通过 user-data 云函数间接读写')
doc.add_paragraph('• 云函数：user-data（核心读写）、admin-api（管理查看/更新）、init-stats（统计聚合）')

doc.add_heading('3.5 典型数据示例', level=2)
doc.add_paragraph('''{
  "_id": "auto_generated_id",
  "_openid": "o4BVT3QcW1wbAgEt1yqRD7drVPhY",
  "favorites": [
    { "routeId": "route_001", "date": "2026-03-30T08:00:00.000Z" },
    { "routeId": "route_005", "date": "2026-03-31T10:30:00.000Z" }
  ],
  "completed": [
    {
      "routeId": "route_003",
      "date": "2026-03-29",
      "name": "翠华山天池环线",
      "weather": "晴天",
      "feeling": "风景很美",
      "difficultyFeeling": "比预期轻松",
      "companions": "老王、小李",
      "distance": 6.5,
      "note": "天池很美，值得再来",
      "completedAt": 1711680000000
    }
  ],
  "checklists": {
    "route_001": { "checked": [0, 1, 3], "custom": ["暖宝宝"] }
  },
  "updatedAt": "2026-03-31T12:00:00.000Z"
}''')


# ===== 4. articles 集合 =====
doc.add_page_break()
doc.add_heading('4. articles（知识文章）', level=1)

doc.add_heading('4.1 集合概述', level=2)
doc.add_paragraph('存储户外知识文章，包括装备选购、安全自救、户外礼仪等分类。在小程序的"知识"页面展示，帮助用户学习户外知识。')

doc.add_heading('4.2 字段结构', level=2)
add_field_table(doc, [
    ('_id', 'String', '文章ID', '文章唯一标识，格式 article_XXX', '是', '系统生成'),
    ('title', 'String', '文章标题', '文章标题', '是', '-'),
    ('category', 'String', '文章分类', '如：装备选购、安全自救、户外礼仪、装备推荐、其他', '是', '-'),
    ('subcategory', 'String', '子分类', '如：必备装备、基础安全、LNT无痕山林', '否', '""'),
    ('difficulty', 'String', '难度级别', 'beginner / intermediate / advanced', '否', '"beginner"'),
    ('priority', 'Number', '优先级', '数值越大优先级越高', '否', '0'),
    ('tags', 'Array[String]', '标签', '文章标签，用于搜索和分类', '否', '[]'),
    ('season', 'Array[String]', '适用季节', '如：["all"] 或 ["spring","autumn"]', '否', '["all"]'),
    ('content', 'String', '文章正文', 'HTML格式的文章内容', '是', '-'),
    ('summary', 'String', '摘要', '文章简短摘要', '否', '""'),
    ('coverImage', 'String', '封面图', '封面图片路径', '否', '""'),
    ('readTime', 'String', '阅读时长', '如：5分钟、6分钟', '否', '""'),
    ('likes', 'Number', '点赞数', '文章获得的点赞数', '否', '0'),
    ('viewCount', 'Number', '阅读次数', '文章被阅读的次数', '否', '0'),
    ('highlights', 'String', '精华亮点', '文章亮点摘要', '否', '""'),
    ('isActive', 'Boolean', '是否启用', '控制文章是否在前台展示', '否', 'true'),
    ('order', 'Number', '排序号', '展示顺序', '否', '999'),
    ('createdAt', 'Date', '创建时间', '文章创建时间', '否', 'db.serverDate()'),
])

doc.add_heading('4.3 索引信息', level=2)
doc.add_paragraph('• order 字段：按升序排列，用于列表默认排序')
doc.add_paragraph('• category 字段：用于按分类筛选文章')
doc.add_paragraph('• isActive 字段：用于筛选有效文章')
doc.add_paragraph('• viewCount 字段：用于阅读量排行聚合')

doc.add_heading('4.4 读写权限', level=2)
doc.add_paragraph('• 前端：通过 articles 云函数查询列表/详情/推荐，阅读次数+1')
doc.add_paragraph('• 云函数：articles（读+viewCount自增）、admin-api（完整CRUD）、import-data（批量导入/清空）、cloud_update_articles（更新）')

doc.add_heading('4.5 典型数据示例', level=2)
doc.add_paragraph('''{
  "_id": "article_050",
  "title": "新手徒步必备装备清单",
  "category": "装备推荐",
  "subcategory": "必备装备",
  "difficulty": "beginner",
  "tags": ["新手", "装备", "清单", "入门"],
  "season": ["all"],
  "content": "<div class=\\"article-body\\"><h2>第一次徒步，带什么？</h2>...</div>",
  "summary": "第一次徒步该带什么？必备装备+建议装备+踩坑清单",
  "coverImage": "/images/scenery/scenery-general.jpg",
  "readTime": "6分钟",
  "likes": 0,
  "viewCount": 0,
  "highlights": "第一次徒步不知道带什么？这份新手装备清单帮你避开90%的坑。",
  "isActive": true,
  "order": 50
}''')


# ===== 5. admin_config 集合 =====
doc.add_page_break()
doc.add_heading('5. admin_config（后台配置）', level=1)

doc.add_heading('5.1 集合概述', level=2)
doc.add_paragraph('存储后台管理系统的配置信息。目前主要存放管理员相关的配置参数，数据量极少（通常只有1条记录）。')

doc.add_heading('5.2 字段结构', level=2)
add_field_table(doc, [
    ('_id', 'String', '文档ID', '系统自动生成', '是', '系统生成'),
    ('createdAt', 'Date', '创建时间', '配置创建时间', '否', 'db.serverDate()'),
    ('(动态字段)', 'Any', '配置项', '后台管理的动态配置参数', '否', '-'),
])

doc.add_heading('5.3 索引信息', level=2)
doc.add_paragraph('• 无特殊索引，通过 limit(1) 获取唯一记录')

doc.add_heading('5.4 读写权限', level=2)
doc.add_paragraph('• 前端：不直接访问')
doc.add_paragraph('• 云函数：admin-api（读取/更新配置）')

doc.add_heading('5.5 典型数据示例', level=2)
doc.add_paragraph('''{
  "_id": "auto_generated_id",
  "createdAt": "2026-03-29T12:00:00.000Z",
  "someConfig": "value"
}''')


# ===== 6. counters 集合 =====
doc.add_page_break()
doc.add_heading('6. counters（计数器）', level=1)

doc.add_heading('6.1 集合概述', level=2)
doc.add_paragraph('存储自增序列计数器。目前主要用于用户编号的自增分配，确保每个用户获得唯一递增的编号。')

doc.add_heading('6.2 字段结构', level=2)
add_field_table(doc, [
    ('_id', 'String', '计数器ID', '计数器名称，如 "user_number"', '是', '"user_number"'),
    ('value', 'Number', '当前值', '当前计数器的值', '是', '0'),
])

doc.add_heading('6.3 索引信息', level=2)
doc.add_paragraph('• _id 字段：通过 doc("user_number") 直接定位唯一记录')

doc.add_heading('6.4 读写权限', level=2)
doc.add_paragraph('• 前端：不直接访问')
doc.add_paragraph('• 云函数：user-data（原子自增操作 _.inc(1)）')

doc.add_heading('6.5 典型数据示例', level=2)
doc.add_paragraph('''{
  "_id": "user_number",
  "value": 42
}''')


# ===== 7. weather 集合 =====
doc.add_page_break()
doc.add_heading('7. weather（天气缓存）', level=1)

doc.add_heading('7.1 集合概述', level=2)
doc.add_paragraph('缓存天气数据，通过定时触发器（每天7:00）自动刷新。支持按城市查询，减少实时API调用。')

doc.add_heading('7.2 字段结构', level=2)
add_field_table(doc, [
    ('_id', 'String', '文档ID', '系统自动生成', '是', '系统生成'),
    ('city', 'String', '城市名称', '查询的城市，如"西安"', '是', '"西安"'),
    ('temp', 'String', '温度', '当前温度（摄氏度）', '否', '"--"'),
    ('desc', 'String', '天气描述', '中文天气描述，如"晴天"、"小雨"', '否', '"暂无数据"'),
    ('icon', 'String', '天气图标', '天气emoji图标', '否', '"🌤️"'),
    ('humidity', 'String', '湿度', '相对湿度百分比', '否', '"--%"'),
    ('wind', 'String', '风力', '风力等级', '否', '"--级"'),
    ('safeTip', 'String', '出行建议', '基于天气的户外安全建议', '否', '""'),
    ('suitable', 'Boolean', '是否适合出行', '根据天气判断是否适合户外活动', '否', 'true'),
    ('updatedAt', 'Date', '更新时间', '最后更新时间', '否', 'db.serverDate()'),
])

doc.add_heading('7.3 索引信息', level=2)
doc.add_paragraph('• city 字段：按城市名称查询缓存数据')

doc.add_heading('7.4 读写权限', level=2)
doc.add_paragraph('• 前端：通过 weather 云函数获取天气数据')
doc.add_paragraph('• 云函数：weather（定时刷新upsert + 实时查询）')

doc.add_heading('7.5 定时触发器', level=2)
doc.add_paragraph('• 触发器名称：daily-weather')
doc.add_paragraph('• 执行时间：每天 07:00（北京时间）')
doc.add_paragraph('• 触发规则：0 0 7 * * * *')

doc.add_heading('7.6 典型数据示例', level=2)
doc.add_paragraph('''{
  "_id": "auto_generated_id",
  "city": "西安",
  "temp": "18",
  "desc": "多云",
  "icon": "⛅",
  "humidity": "65%",
  "wind": "2级",
  "safeTip": "天气适宜，适合户外活动",
  "suitable": true,
  "updatedAt": "2026-04-03T07:00:00.000Z"
}''')


# ===== 数据关系说明 =====
doc.add_page_break()
doc.add_heading('数据关系说明', level=1)

doc.add_heading('集合间关系图', level=2)
doc.add_paragraph('''
┌─────────────┐     ┌──────────────┐     ┌─────────────┐
│   counters   │     │    users     │     │   weather   │
│ (user_number)│     │  (用户信息)   │     │  (天气缓存)  │
└─────────────┘     └──────┬───────┘     └─────────────┘
                           │ _openid
                           │
                    ┌──────┴───────┐
                    │  user_data   │
                    │  (用户数据)   │
                    └──┬────────┬──┘
                       │        │
          favorites[].routeId   completed[].routeId
                       │        │
                    ┌──┴────────┴──┐
                    │    routes    │
                    │  (徒步路线)   │◄──── favoriteCount / completedCount
                    └──────────────┘

┌──────────────┐     ┌────────────────┐
│   articles   │     │  admin_config  │
│  (知识文章)   │     │  (后台配置)     │
└──────────────┘     └────────────────┘
''')

doc.add_heading('核心关系说明', level=2)

doc.add_heading('users ↔ user_data（一对多→一对一）', level=3)
doc.add_paragraph('• 通过 _openid 字段关联')
doc.add_paragraph('• users 记录用户基本信息（编号、昵称、访问次数）')
doc.add_paragraph('• user_data 记录用户的行为数据（收藏、已走过、清单）')
doc.add_paragraph('• 一个用户在 users 中有一条记录，在 user_data 中也最多有一条记录')

doc.add_heading('user_data → routes（多对多）', level=3)
doc.add_paragraph('• favorites[].routeId 引用 routes._id（收藏关系）')
doc.add_paragraph('• completed[].routeId 引用 routes._id（已走过关系）')
doc.add_paragraph('• user_data 中使用 routeId 冗余存储路线信息')
doc.add_paragraph('• routes 集合通过 favoriteCount 和 completedCount 反向统计')

doc.add_heading('counters → users（编号分配）', level=3)
doc.add_paragraph('• counters 集合中的 "user_number" 文档维护全局自增值')
doc.add_paragraph('• 新用户注册时，原子自增 counters.value，分配给 users.userNumber')

doc.add_heading('weather（独立缓存）', level=3)
doc.add_paragraph('• weather 集合独立运行，通过定时触发器自动刷新')
doc.add_paragraph('• 按城市名（city）作为唯一键进行 upsert 操作')

doc.add_heading('数据一致性要点', level=2)
doc.add_paragraph('• 收藏/取消收藏时，同时更新 user_data.favorites 和 routes.favoriteCount（原子操作 _.inc）')
doc.add_paragraph('• 添加/删除已走过记录时，同时更新 user_data.completed 和 routes.completedCount')
doc.add_paragraph('• 用户编号通过 counters 集合的原子自增（_.inc(1)）保证唯一性')
doc.add_paragraph('• 天气数据通过 upsert（先查后插/更新）保证城市唯一性')


# ===== 云函数与集合对应关系 =====
doc.add_heading('云函数与集合对应关系', level=2)

func_table = doc.add_table(rows=8, cols=3, style='Light Grid Accent 1')
func_table.rows[0].cells[0].text = '云函数'
func_table.rows[0].cells[1].text = '操作的集合'
func_table.rows[0].cells[2].text = '主要功能'
func_data = [
    ('routes', 'routes', '路线列表/详情/搜索查询'),
    ('articles', 'articles', '文章列表/详情/推荐/阅读计数'),
    ('user-data', 'user_data, users, counters, routes', '收藏/已走过/清单管理、用户注册'),
    ('weather', 'weather', '天气查询/定时缓存刷新'),
    ('admin-api', 'routes, users, user_data, articles, admin_config', '后台管理：CRUD、统计、导出'),
    ('import-data', 'routes, articles', '批量导入/清空/更新数据'),
    ('init-stats', 'routes, user_data', '初始化路线收藏/已走过次数统计'),
]
for i, (f, c, d) in enumerate(func_data):
    func_table.rows[i+1].cells[0].text = f
    func_table.rows[i+1].cells[1].text = c
    func_table.rows[i+1].cells[2].text = d


# ===== 保存 =====
output_path = '/Users/wangweixin/.openclaw/media/outbound/秦人户外小程序数据库说明文档.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'文档已生成：{output_path}')
